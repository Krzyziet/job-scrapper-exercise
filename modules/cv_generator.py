import os
import re
import logging
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path("generated_cvs")

# Sekcje CV dopasowane do trybu podkreślenia
CV_SECTIONS = {
    "network": {
        "summary": (
            "Doświadczony Technical Product Owner z backgroundem Network Engineera. "
            "Zarządzałem 18-osobowym zespołem Network Operations w ING Hubs Poland. "
            "Specjalizuję się w ITSM, zarządzaniu infrastrukturą sieciową (Cisco, Palo Alto, F5) "
            "oraz procesach Incident & Change Management w środowiskach krytycznych."
        ),
        "highlights": [
            "Zarządzanie 18-osobowym zespołem Network Operations – ING Hubs Poland",
            "Wdrożenie i utrzymanie infrastruktury sieciowej: Cisco, Palo Alto, F5, VMware ESXi",
            "ITSM: Incident Management, Change Management, Problem Management (HP Service Manager)",
            "Certyfikaty: CCNP Enterprise, CCNA, PSPO I",
        ],
    },
    "management": {
        "summary": (
            "Technical Product Owner i lider zespołu z 5+ letnim doświadczeniem w dużych korporacjach. "
            "Zarządzałem 18-osobowym zespołem operacyjnym w ING Hubs Poland, nagrodzony Director's Award. "
            "Łączę kompetencje techniczne z umiejętnościami zarządzania produktem, roadmapą i stakeholderami."
        ),
        "highlights": [
            "Zarządzanie 18-osobowym zespołem – Director's Award za wyniki",
            "Product ownership: roadmap, backlog, sprint planning, stakeholder management",
            "Certyfikat PSPO I (Professional Scrum Product Owner)",
            "5+ lat w środowisku bankowym (ING Hubs Poland)",
        ],
    },
    "product": {
        "summary": (
            "Technical Product Owner z silnym backgroundem technicznym i doświadczeniem agile. "
            "Posiadam certyfikat PSPO I. Zarządzam backlogiem, roadmapą produktową i relacjami "
            "ze stakeholderami w środowisku bankowym. Łączę perspektywę techniczną z business value."
        ),
        "highlights": [
            "Certyfikat PSPO I – Scrum.org",
            "Roadmap planning, backlog refinement, sprint reviews",
            "Stakeholder management na poziomie C-level",
            "Doświadczenie w bankowości i środowiskach regulated",
        ],
    },
    "devops": {
        "summary": (
            "Technical Product Owner z doświadczeniem w automatyzacji i środowiskach cloud/wirtualizacji. "
            "Pracowałem z VMware ESXi, wdrażałem procesy CI/CD w środowiskach sieciowych. "
            "Łączę kompetencje inżynierskie z product ownershipem w dużych organizacjach IT."
        ),
        "highlights": [
            "VMware ESXi – wirtualizacja infrastruktury",
            "Automatyzacja procesów operacyjnych w środowisku NOC",
            "Doświadczenie z CI/CD i narzędziami DevOps",
            "Certyfikaty: CCNP Enterprise, PSPO I",
        ],
    },
}

EXPERIENCE = [
    {
        "role": "Technical Product Owner",
        "company": "ING Hubs Poland",
        "period": "2019 – obecnie",
        "bullets": [
            "Product ownership dla platformy Network Operations – roadmap, backlog, OKR",
            "Zarządzanie 18-osobowym zespołem inżynierów sieciowych",
            "Stakeholder management: dyrekcja IT, compliance, dostawcy zewnętrzni",
            "Wdrożenie procesów ITSM: Incident, Change, Problem Management",
            "Nagroda Director's Award za transformację operacyjną zespołu",
        ],
    },
    {
        "role": "Network Engineer",
        "company": "ING Hubs Poland",
        "period": "2017 – 2019",
        "bullets": [
            "Projektowanie i utrzymanie infrastruktury sieciowej (Cisco, Palo Alto, F5)",
            "Zarządzanie firewallami i load balancerami w środowisku bankowym",
            "Wirtualizacja: VMware ESXi, konfiguracja vSwitchów",
            "On-call support dla systemów krytycznych 24/7",
        ],
    },
]

EDUCATION = [
    {"degree": "Inżynier – Teleinformatyka", "school": "Politechnika Łódzka", "year": "2015"},
]

CERTS = ["PSPO I – Scrum.org", "CCNP Enterprise – Cisco", "CCNA – Cisco"]


def _add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)  # niebieski ING-style


def _add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text).font.size = Pt(10)


def generate_cv(offer: dict) -> Path:
    """Generuje plik .docx z CV dopasowanym do oferty. Zwraca ścieżkę pliku."""
    OUTPUT_DIR.mkdir(exist_ok=True)

    emphasis = offer.get("cv_emphasis", "management")
    section = CV_SECTIONS.get(emphasis, CV_SECTIONS["management"])

    company_slug = re.sub(r"[^\w]", "_", offer.get("company", "firma"))[:20]
    title_slug = re.sub(r"[^\w]", "_", offer.get("title", "rola"))[:20]
    filename = OUTPUT_DIR / f"CV_KrzysztofZielinski_{company_slug}_{title_slug}.docx"

    doc = Document()

    # Nagłówek
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_para.add_run("Krzysztof Zieliński")
    run.bold = True
    run.font.size = Pt(18)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.add_run(
        "krzy.ziet@gmail.com  |  Łódź, Polska  |  LinkedIn  |  +48 XXX XXX XXX"
    ).font.size = Pt(9)

    doc.add_paragraph()

    # Podsumowanie
    _add_heading(doc, "Podsumowanie zawodowe")
    p = doc.add_paragraph(section["summary"])
    p.runs[0].font.size = Pt(10)

    # Kluczowe osiągnięcia (dopasowane)
    _add_heading(doc, "Kluczowe kompetencje")
    for item in section["highlights"]:
        _add_bullet(doc, item)

    # Doświadczenie
    _add_heading(doc, "Doświadczenie zawodowe")
    for exp in EXPERIENCE:
        p = doc.add_paragraph()
        r = p.add_run(f"{exp['role']}  –  {exp['company']}")
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(f"  ({exp['period']})").font.size = Pt(9)
        for bullet in exp["bullets"]:
            _add_bullet(doc, bullet)
        doc.add_paragraph()

    # Wykształcenie
    _add_heading(doc, "Wykształcenie")
    for edu in EDUCATION:
        p = doc.add_paragraph()
        p.add_run(f"{edu['degree']}, {edu['school']}, {edu['year']}").font.size = Pt(10)

    # Certyfikaty
    _add_heading(doc, "Certyfikaty")
    for cert in CERTS:
        _add_bullet(doc, cert)

    # Stopka z notą do oferty
    if offer.get("cover_note"):
        doc.add_paragraph()
        _add_heading(doc, "Dlaczego ta rola")
        p = doc.add_paragraph(offer["cover_note"])
        p.runs[0].font.size = Pt(10)

    doc.save(filename)
    logger.info(f"[CV] wygenerowano: {filename}")
    return filename


def generate_all_cvs(offers: list[dict]) -> list[dict]:
    """Generuje CV dla każdej oferty, dodaje pole 'cv_path' do słownika."""
    for offer in offers:
        try:
            path = generate_cv(offer)
            offer["cv_path"] = str(path)
        except Exception as e:
            logger.error(f"[CV] błąd dla {offer.get('title')}: {e}")
            offer["cv_path"] = ""
    return offers
